# -*- coding: utf-8 -*-
# SPDX-License-Identifier: GPL-3.0-or-later
"""以 .bak 原始指令文件为基底，定向更新指令格式，保留原版式。
- 猎人：合并静止卡两条、删除护盾卡(获线索)错误条目、更新占位字母、补冷却说明
- 任务点：补“全局免疫 15 秒”说明
结果写入对应的非 .bak docx（原 .bak 不动，作为纯净备份）。
"""
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DATA = Path(r"D:/58系统/Claw/data")


def set_runs(p, segments):
    """清空 p 的 runs，按 segments=[(text,bold),...] 重建。"""
    for r in list(p._p.findall(qn("w:r"))):
        r.getparent().remove(r)
    for text, bold in segments:
        r = p.add_run(text)
        r.font.bold = bold


def insert_note_after(ref_p, segments):
    new_p = OxmlElement("w:p")
    ref_p._p.addnext(new_p)
    p = Paragraph(new_p, ref_p._parent)
    for text, bold in segments:
        r = p.add_run(text)
        r.font.bold = bold
    return p


def find_para(doc, text):
    for p in doc.paragraphs:
        if p.text == text:
            return p
    raise ValueError("not found: " + text)


def delete_para(p):
    p._element.getparent().remove(p._element)


# ===================== 猎人指令公式 =====================
src = DATA / "猎人指令公式.bak_before_cmdupdate.docx"
dst = DATA / "猎人指令公式.docx"
shutil.copy(src, dst)
doc = Document(dst)

# 更新指令行文本（保留 Consolas 字体，仅改 .text）
find_para(doc, "X组YYY被ZZZ淘汰").runs[0].text = "X组XXX被XXX淘汰"
find_para(doc, "X组YYY被ZZZ在W淘汰 召唤机动人员").runs[0].text = "X组XXX被XXX在XXX淘汰 召唤机动人员"
find_para(doc, "3. 使用静止卡（不获得线索）").runs[0].text = "3. 使用静止卡（猎人送出线索 / 露水给玩家）"
find_para(doc, "猎人M被使用静止卡").runs[0].text = "猎人Y被使用静止卡 [获得线索N] [获得露水M]"
find_para(doc, "5. 使用护盾卡（不获得线索）").runs[0].text = "4. 使用护盾卡（不获得线索）"
find_para(doc, "猎人M被使用护盾卡").runs[0].text = "猎人Y被使用护盾卡"

# 展开静止卡说明（猎人送出语义 + 180s 冷却）
set_runs(find_para(doc, "触发 180 秒冷却。"), [
    ('线索 / 露水可同条、顺序任意；猎人须先经网页持有该编号，送出后置"已发现未收集"并消耗持有。触发 ', False),
    ("180 秒", True),
    (" 冷却。", False),
])

# 在两条淘汰指令后补冷却说明
insert_note_after(
    find_para(doc, "X组XXX被XXX淘汰"),
    [("触发 ", False), ("20 秒", True), (" 冷却（狂欢 40 秒）。", False)],
)
insert_note_after(
    find_para(doc, "X组XXX被XXX在XXX淘汰 召唤机动人员"),
    [("触发 ", False), ("20 秒", True), (" 冷却（狂欢 40 秒）。", False)],
)

# 删除过时/错误条目：静止卡(获线索) 与 护盾卡(获线索)
for t in [
    "4. 使用静止卡（同时获得线索）",
    "猎人M被使用静止卡 获得线索N",
    '线索 N 标记"已发现未收集" + 180 秒冷却。',
    "6. 使用护盾卡（同时获得线索）",
    "猎人M被使用护盾卡 获得线索N",
    '线索 N 标记"已发现未收集" + 20 秒冷却。',
]:
    delete_para(find_para(doc, t))

doc.save(dst)
print("SAVED", dst)


# ===================== 任务点NPC指令公式 =====================
src2 = DATA / "任务点NPC指令公式.bak_before_cmdupdate.docx"
dst2 = DATA / "任务点NPC指令公式.docx"
shutil.copy(src2, dst2)
doc2 = Document(dst2)

# 任务点完成：补全局免疫 15 秒说明
insert_note_after(
    find_para(doc2, "任务点T 完成 线索C 露水D 金露水G 护盾卡H 静止卡S"),
    [("送出道具并触发", False), ("全局免疫 15 秒", True), ("（全体玩家不可被淘汰）。", False)],
)

doc2.save(dst2)
print("SAVED", dst2)
print("ALL DONE")
