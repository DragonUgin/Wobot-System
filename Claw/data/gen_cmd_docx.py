# -*- coding: utf-8 -*-
# SPDX-License-Identifier: GPL-3.0-or-later
"""重新生成「猎人指令公式.docx」与「任务点NPC指令公式.docx」。

占位符统一使用大写字母（X/Y/Z/N/M/XXX），与 src/plugins/work_group.py 文档字符串一致。
运行：python gen_cmd_docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).parent

AMBER = RGBColor(0xD9, 0x77, 0x06)
GRAY = RGBColor(0x6B, 0x72, 0x80)
RED = RGBColor(0xDC, 0x26, 0x26)
BLACK = RGBColor(0x1F, 0x1F, 0x1F)


def _set_font(run, size=11, bold=False, color=BLACK, mono=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Consolas" if mono else "Microsoft YaHei"


def add_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=16, bold=True, color=AMBER)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=13, bold=True, color=BLACK)
    return p


def add_note(doc, text, color=GRAY, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=size, color=color)
    return p


def add_cmd(doc, fmt_text):
    """指令格式行（等宽字体）。"""
    p = doc.add_paragraph()
    r = p.add_run(fmt_text)
    _set_font(r, size=12, bold=True, color=BLACK, mono=True)
    return p


def add_example(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("示例：")
    _set_font(r, size=10, color=GRAY)
    r2 = p.add_run(text)
    _set_font(r2, size=10, color=GRAY, mono=True)
    return p


def add_effect(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("效果：")
    _set_font(r, size=10, color=GRAY)
    r2 = p.add_run(text)
    _set_font(r2, size=10, color=GRAY)
    return p


def add_legend(doc, rows):
    """占位符图例：rows = [(字母, 含义), ...]"""
    add_subheading(doc, "占位符图例")
    for letter, meaning in rows:
        p = doc.add_paragraph()
        r = p.add_run(letter)
        _set_font(r, size=11, bold=True, color=AMBER, mono=True)
        r2 = p.add_run("  " + meaning)
        _set_font(r2, size=11, color=BLACK)


def add_rule(doc, text):
    """注意事项条目。"""
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_font(r, size=10, color=GRAY)
    return p


# ===== 猎人指令公式 =====
def build_hunter_doc():
    doc = Document()
    add_heading(doc, "猎人指令公式")
    add_note(doc, "所有指令在工作群发送，仅游戏运行中生效。猎人身份 = 工作群成员（去名单化）。", color=GRAY)

    add_legend(doc, [
        ("X", "组号（数字）"),
        ("Y", "猎人名（静止卡 / 护盾卡）"),
        ("N", "线索编号（静止卡送出）"),
        ("M", "露水编号（静止卡送出）"),
        ("XXX", "玩家名 / 猎人名 / 地点（多字符占位）"),
    ])

    add_subheading(doc, "1. 淘汰（无道具）")
    add_cmd(doc, "X组XXX被XXX淘汰")
    add_example(doc, "1组张三被李四淘汰")
    add_effect(doc, "淘汰玩家；猎人冷却 20 秒（狂欢 40 秒）；全局免疫期不可用。")

    add_subheading(doc, "2. 淘汰（带道具护送）")
    add_cmd(doc, "X组XXX被XXX在XXX淘汰 召唤机动人员")
    add_example(doc, "1组张三被李四在学校淘汰 召唤机动人员")
    add_effect(doc, "淘汰并记录地点，通知后台群机动护送；冷却同上。")

    add_subheading(doc, "3. 静止卡（唯一可让玩家获得线索+露水的卡）")
    add_cmd(doc, "猎人Y被使用静止卡 [获得线索N] [获得露水M]")
    add_example(doc, "猎人李四被使用静止卡 获得线索2 获得露水501")
    add_effect(doc, "猎人送出线索/露水给玩家（消耗其持有），置「已发现未收集」；冷却 180 秒。")
    add_note(doc, "说明：「获得线索N」「获得露水M」各自独立子句，可只写一个、可同条出现、顺序任意。"
                  "猎人须先经网页分配持有该编号，否则告警不送出。", color=GRAY)

    add_subheading(doc, "4. 护盾卡")
    add_cmd(doc, "猎人Y被使用护盾卡")
    add_example(doc, "猎人李四被使用护盾卡")
    add_effect(doc, "确认使用护盾卡；冷却 20 秒。不再支持获得线索。")

    add_subheading(doc, "注意事项")
    add_rule(doc, "淘汰冷却按发送者 QQ 记录；静止卡 / 护盾卡冷却按猎人名记录。")
    add_rule(doc, "全局任务免疫期间（任意任务点完成后 15 秒）所有淘汰指令被拦截。")
    add_rule(doc, "猎人持有物（线索 / 露水）只能由管理员网页导入 / 分配，游戏指令只消耗不新增。")
    add_rule(doc, "已废弃指令：线索N 已收集（改由静止卡 / 任务点完成触发）。")

    out = OUT_DIR / "猎人指令公式.docx"
    doc.save(out)
    print("SAVED", out)


# ===== 任务点NPC指令公式 =====
def build_tp_doc():
    doc = Document()
    add_heading(doc, "任务点NPC指令公式")
    add_note(doc, "所有指令在工作群发送，仅游戏运行中生效。任务点 NPC 身份 = 工作群成员（去名单化）。", color=GRAY)

    add_legend(doc, [
        ("N", "任务点编号"),
        ("X", "线索编号（填编号或「无」）"),
        ("Y", "露水数量（填数量或「无」）"),
        ("Z", "金露水 / 护盾卡 / 静止卡 数量（填数量或「无」）"),
    ])

    add_subheading(doc, "1. 任务点完成（送出道具）")
    add_cmd(doc, "任务点N 完成 线索X 露水Y 金露水Z 护盾卡Z 静止卡Z")
    add_example(doc, "任务点3 完成 线索5 露水2 金露水无 护盾卡无 静止卡无")
    add_effect(doc, "送出道具给玩家；触发全局免疫 15 秒（全体玩家不可被淘汰）。")

    add_subheading(doc, "2. 任务点接收（入库道具）")
    add_cmd(doc, "任务点N 接收 线索X 露水Y 金露水Z 护盾卡Z 静止卡Z")
    add_example(doc, "任务点3 接收 线索无 露水4 金露水1 护盾卡2 静止卡无")
    add_effect(doc, "道具入库；不触发免疫。")

    add_subheading(doc, "注意事项")
    add_rule(doc, "指令须精确匹配：关键字「完成 / 接收 / 线索 / 露水 / 金露水 / 护盾卡 / 静止卡」必须原样，多余空格可容忍。")
    add_rule(doc, "编号位填数字，数量位填数字或「无」。")
    add_rule(doc, "任务点完成后 15 秒内全体玩家免疫淘汰（全局任务免疫）。")
    add_rule(doc, "任务点 NPC 名单已去名单化，靠工作群成员资格保证身份，系统只验证业务指令格式。")

    out = OUT_DIR / "任务点NPC指令公式.docx"
    doc.save(out)
    print("SAVED", out)


if __name__ == "__main__":
    build_hunter_doc()
    build_tp_doc()
    print("ALL DONE")
