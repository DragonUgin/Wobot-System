# -*- coding: utf-8 -*-
# SPDX-License-Identifier: GPL-3.0-or-later
"""把两张流程图插入到对应 docx 的"附录：...流程图"标题下方。"""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from pathlib import Path

DATA = Path(r"D:/58系统/Claw/data")
IMG_DIR = Path(r"D:/Program Files (x86)/weixinDownload/xwechat_files/wxid_1370353703312_ac1b/temp/RWTemp/2026-08/7c3920fc350133b544113faebfb50898")


def insert_image_after_heading(doc_path, heading_text, image_path, width_inches=6.0):
    doc = Document(doc_path)
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == heading_text:
            target = p
            break
    if target is None:
        raise ValueError(f"未找到标题: {heading_text}")

    new_p = OxmlElement("w:p")
    target._p.addnext(new_p)
    para = Paragraph(new_p, target._parent)
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    doc.save(doc_path)
    print(f"已在 {doc_path.name} 的「{heading_text}」后插入 {image_path.name}")


insert_image_after_heading(
    DATA / "猎人指令公式.docx",
    "附录：猎人指令流程图",
    IMG_DIR / "bd5299575cc89968e3cf9430f25afb38.jpg",
)

insert_image_after_heading(
    DATA / "任务点NPC指令公式.docx",
    "附录：任务点NPC指令流程图",
    IMG_DIR / "796b816e1f9c2c6d883eff559d90b44a.jpg",
)

print("ALL DONE")
